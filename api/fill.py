from docx import Document
from docx.oxml import OxmlElement
from models import SubmissionParams
from docx.shared import Inches
from docx.oxml.shared import qn


def gen_filled_doc(data: SubmissionParams, doc: Document):
    replacements = {
        "submission_type": data.submission_type,
        "topic_name": data.topic_name,
        "subject_name": data.subject_name,
        "subject_code": data.subject_code.upper(),
        "semester_number": str(data.semester_number),
        "student_branch": data.student_branch.title(),
        "faculty_branch": data.faculty_branch.title(),
        "faculty_name": data.faculty_name_with_title.title(),
        "designation": data.designation.title(),
        "from_ay": str(data.from_ay),
        "to_ay": str(data.to_ay),
    }

    for paragraph in doc.paragraphs:
        for placeholder, replacement in replacements.items():
            replace_text_robust(paragraph, placeholder, replacement)

        process_submitters_robust(paragraph, data.submitters)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for placeholder, replacement in replacements.items():
                        replace_text_robust(paragraph, placeholder, replacement)
                    process_submitters_robust(paragraph, data.submitters)

    add_page_border_and_margins(doc)

    temp_docx_path = "frontpage.docx"
    doc.save(temp_docx_path)
    return temp_docx_path


def replace_text_robust(paragraph, placeholder, replacement):
    if placeholder in paragraph.text:
        full_text = paragraph.text
        if placeholder in full_text:
            new_text = full_text.replace(placeholder, replacement)

            for run in paragraph.runs:
                run.text = ""

            if paragraph.runs:
                paragraph.runs[0].text = new_text
            else:
                paragraph.add_run(new_text)


def process_submitters_robust(paragraph, submitters):
    for i in range(1, 8):
        placeholder = f"submitter{i}"
        if placeholder in paragraph.text:
            if i <= len(submitters):
                name, usn = list(submitters.items())[i - 1]
                replacement = f"{name.title()} ({usn.upper()})"
                replace_text_robust(paragraph, placeholder, replacement)
            else:
                replace_text_robust(paragraph, placeholder, "")


def add_page_border_and_margins(doc):
    section = doc.sections[0]

    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    sectPr = section._sectPr
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")

    border_attrs = {"val": "single", "sz": "20", "space": "24", "color": "000000"}

    for side in ["top", "left", "bottom", "right"]:
        border = OxmlElement(f"w:{side}")

        for attr, value in border_attrs.items():
            border.set(qn(f"w:{attr}"), value)

        pgBorders.append(border)

    sectPr.append(pgBorders)
